import os
import re
import pytesseract
import fitz
import uuid
import PyPDF2
from functools import lru_cache
from datetime import datetime
from typing import List, Dict, Any

from docx import Document
from PIL import Image, ImageOps, ImageEnhance
from collections import Counter
from langdetect import detect
from typing import Dict, Any, List
from datetime import datetime


VECTORIZER_MODEL = None
STOPWORD_COLLECTION = None


class PreprocessingService:
    def __init__(self, stopwords_collection: set, vectorizer_model: object) -> None:
        self.stopwords = stopwords_collection
        self.vectorizer_model = vectorizer_model
        
    """
    Извлечение и генерация метаданных
    """
    def parse_pdf_date(self, date_str: str) -> datetime:
        """
        Парсит дату из PDF метаданных в различных форматах.
        
        Args:
            date_str: Строка с датой в формате PDF (например, "D:20230512163128+03'00'")
        
        Returns:
            datetime: Объект datetime с датой создания
        """
        # Убираем префикс 'D:' если есть
        date_str = date_str.replace('D:', '')
        
        # Извлекаем основную часть даты (YYYYMMDDHHMMSS)
        date_pattern = r'(\d{14}|\d{12}|\d{8})'
        match = re.search(date_pattern, date_str)
        
        if not match:
            raise ValueError(f"Неверный формат даты: {date_str}")
            
        date_str = match.group(1)
        
        # Обрабатываем разные длины дат
        if len(date_str) == 14:  # YYYYMMDDHHMMSS
            return datetime.strptime(date_str, '%Y%m%d%H%M%S')
        elif len(date_str) == 12:  # YYYYMMDDHHMM
            return datetime.strptime(date_str, '%Y%m%d%H%M')
        else:  # YYYYMMDD
            return datetime.strptime(date_str, '%Y%m%d')

    def extract_metadata_pdf(self, pdf_path: str) -> Dict[str, any]:
        """
        Извлекает автора и дату создания из PDF-документа.
        Возвращает словарь с ключами 'author' и 'created_date'.
        """
        try:
            metadata = {'author': None, 'created_date': None}
            
            with open(pdf_path, 'rb') as f:
                pdf = PyPDF2.PdfReader(f)
                
                # Проверяем, зашифрован ли документ
                if pdf.is_encrypted:
                    try:
                        pdf.decrypt('')
                    except:
                        print(f"PDF файл {pdf_path} защищен паролем")
                        return metadata
                
                if pdf.metadata:
                    # Извлекаем автора
                    author = pdf.metadata.get('/Author', None)
                    if author:
                        metadata['author'] = str(author)

                    # Извлекаем дату создания
                    created_date = pdf.metadata.get('/CreationDate', None)
                    if created_date:
                        try:
                            date_obj = self.parse_pdf_date(str(created_date))
                            metadata['created_date'] = date_obj.strftime('%Y-%m-%d')
                        except Exception as e:
                            print(f"Не удалось преобразовать дату создания: {created_date}. Ошибка: {str(e)}")
                            
                            # Пробуем альтернативное поле даты, если есть
                            mod_date = pdf.metadata.get('/ModDate', None)
                            if mod_date:
                                try:
                                    date_obj = self.parse_pdf_date(str(mod_date))
                                    metadata['created_date'] = date_obj.strftime('%Y-%m-%d')
                                except:
                                    print(f"Не удалось преобразовать альтернативную дату: {mod_date}")
                    
            return metadata
        except Exception as e:
            print(f"Не удалось извлечь метаданные из PDF: {e}")
            return {'author': None, 'created_date': None}

    def extract_metadata_word(self, docx_path: str) -> Dict[str, any]:
        """
        Извлекает автора и дату создания из Word-документа.
        Возвращает словарь с ключами 'author' и 'created_date'.
        """
        try:
            metadata = {'author': None, 'created_date': None}
            
            doc = Document(docx_path)
            core_props = doc.core_properties
            
            # Извлекаем автора
            author = core_props.author
            if author:
                metadata['author'] = str(author)
                
            # Извлекаем дату создания
            created_date = core_props.created
            if created_date:
                metadata['created_date'] = created_date.strftime('%Y-%m-%d')
                
            return metadata
        except Exception as e:
            print(f"Не удалось извлечь метаданные из Word: {e}")
            return {'author': None, 'created_date': None}

    def preprocess_text(self, text: str) -> List[str]:
        """
        Очищает текст: удаляет пунктуацию, разрезает на слова и приводит к нижнему регистру.
        """
        # Убираем пунктуацию и лишние символы
        text = re.sub(r'[^\w\s]', '', text)
        # Разбиваем на слова и приводим к нижнему регистру
        words = text.lower().split()
        return words

    def generate_tags_multilang(self, text: str, num_tags: int = 5) -> List[str]:
        """
        Генерирует ключевые слова на основе текста, поддерживает русский и английский языки.
        
        Args:
            text (str): Текст, из которого извлекаем ключевые слова.
            num_tags (int): Количество выводимых ключевых слов.

        Returns:
            List[str]: Список ключевых слов.
        """
        try:
            # Определяем язык текста
            lang = detect(text)
            print(f"Определён язык текста: {lang}")

            # Выбираем стоп-слова в зависимости от языка
            if lang == 'ru':
                stop_words = self.stopwords
            else:
                stop_words = self.stopwords

            # Токенизация текста
            words = self.preprocess_text(text)

            # Убираем стоп-слова
            meaningful_words = [word for word in words if word not in stop_words]

            # Подсчитываем частоту слов
            word_counts = Counter(meaningful_words)

            # Извлекаем num_tags самых частых слов
            tags = [word for word, _ in word_counts.most_common(num_tags)]
            
            return tags

        except Exception as e:
            print(f"Ошибка формирования тегов: {e}")
            return []
        
    """
    Извлечение текста
    """
    def extract_text_from_pdf(self, pdf_path: str) -> str:
        """Извлекает текст из PDF-документа, включая зашифрованные файлы."""
        try:
            text = ""
            with open(pdf_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                
                # Проверяем, зашифрован ли документ
                if reader.is_encrypted:
                    try:
                        # Пробуем расшифровать без пароля
                        reader.decrypt('')
                    except:
                        print(f"PDF файл {pdf_path} защищен паролем")
                        return ""
                
                # Извлекаем текст
                for page in reader.pages:
                    try:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text
                    except Exception as e:
                        print(f"Ошибка при извлечении текста со страницы: {e}")
                        continue
                        
            return text.strip()
        except Exception as e:
            print(f"Ошибка извлечения текста из PDF {pdf_path}: {e}")
            return ""

    def extract_text_from_word(self, word_path: str) -> str:
        """Извлекает текст из Word-документа."""
        try:
            doc = Document(word_path)
            return "\n".join(paragraph.text for paragraph in doc.paragraphs).strip()
        except Exception as e:
            print(f"Ошибка извлечения текста из Word {word_path}: {e}")
            return ""
        
    """ 
    Извлечение изображений
    """
    def create_document_folder(self, document_id: str) -> str:
        """Создает папку для хранения изображений документа."""
        try:
            folder_path = f"data/{document_id}"
            os.makedirs(folder_path, exist_ok=True)
            return folder_path
        except Exception as e:
            print(f"Ошибка создания папки: {e}")
            return ""

    def preprocess_image(self, image_path: str) -> Image:

        """Комплексная предобработка изображения."""
        
        try:
        
            # Загрузка изображения
        
            image = Image.open(image_path)

            # Конвертация в оттенки серого
        
            gray_image = ImageOps.grayscale(image)
        
            # Увеличение резкости
        
            enhancer = ImageEnhance.Sharpness(gray_image)
        
            final_image = enhancer.enhance(2.0)

            return final_image
        
        except Exception as e:
        
            print(f"Ошибка предобработки изображения: {e}")
        
            return Image.open(image_path)

    def has_text_in_image(self, image_path: str) -> bool:
        """Проверяет, содержит ли изображение текст."""
        try:
            # Предобработка изображения
            processed_image = self.preprocess_image(image_path)
            # Конфигурация OCR
            config = r'--oem 3 --psm 3'
            # Распознавание текста с улучшенной конфигурацией
            text = pytesseract.image_to_string(processed_image, lang="rus+eng", config=config)
            return bool(text.strip())
        except Exception as e:
            print(f"Ошибка проверки текста в изображении {image_path}: {e}")
            return False

    def extract_images_from_pdf(self, pdf_path: str, folder_path: str) -> List[str]:
        """Извлекает изображения из PDF-документа и сохраняет только те, что содержат текст."""
        images = []
        try:
            doc = fitz.open(pdf_path)
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                for img_index, img in enumerate(page.get_images(full=True)):
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    image_bytes = base_image["image"]
                    image_ext = base_image["ext"]
                    # Временно сохраняем изображение для проверки
                    temp_path = f"{folder_path}/temp_image.{image_ext}"
                    with open(temp_path, "wb") as f:
                        f.write(image_bytes)
                    # Проверяем наличие текста
                    if self.has_text_in_image(temp_path):
                        # Если текст есть, сохраняем изображение с нормальным именем
                        image_path = f"{folder_path}/image_page{page_num + 1}_img{img_index + 1}.{image_ext}"
                        os.rename(temp_path, image_path)
                        images.append(image_path)
                    else:
                        # Если текста нет, удаляем временный файл
                        os.remove(temp_path)
        except Exception as e:
            print(f"Ошибка извлечения изображений из PDF {pdf_path}: {e}")
        return images

    def extract_images_from_word(self, word_path: str, folder_path: str) -> List[str]:
        """Извлекает изображения из Word-документа и сохраняет только те, что содержат текст."""
        images = []
        try:
            doc = Document(word_path)
            for index, rel in enumerate(doc.part.rels.values()):
                if "image" in rel.target_ref:
                    image_data = rel.target_part.blob
                    # Временно сохраняем изображение для проверки
                    temp_path = f"{folder_path}/temp_image_{index}.png"
                    with open(temp_path, "wb") as f:
                        f.write(image_data)
                    # Проверяем наличие текста
                    if self.has_text_in_image(temp_path):
                        # Если текст есть, сохраняем изображение с нормальным именем
                        image_path = f"{folder_path}/image_word_{len(images) + 1}.png"
                        os.rename(temp_path, image_path)
                        images.append(image_path)
                    else:
                        # Если текста нет, удаляем временный файл
                        os.remove(temp_path)
        except Exception as e:
            print(f"Ошибка извлечения изображений из Word {word_path}: {e}")
        return images

    """
    OCR для изображений
    """
    def clean_ocr_text(self, text: str) -> str:
        """Очищает распознанный текст от мусора и нормализует его."""
        try:
            # Удаляем специальные символы, оставляя буквы, цифры, точки и запятые
            cleaned = re.sub(r'[^a-zA-Zа-яА-Я0-9.,\s]', '', text)
            
            # Заменяем множественные пробелы на один
            cleaned = re.sub(r'\s+', ' ', cleaned)
            
            # Удаляем пробелы перед знаками препинания
            cleaned = re.sub(r'\s+([.,])', r'\1', cleaned)
            
            # Удаляем отдельно стоящие буквы (вероятные ошибки распознавания)
            cleaned = re.sub(r'\s+[a-zA-Zа-яА-Я]\s+', ' ', cleaned)
            
            # Удаляем строки, содержащие менее 2 символов
            lines = [line.strip() for line in cleaned.split('\n') if len(line.strip()) > 2]
            
            return '\n'.join(lines).strip()
        except Exception as e:
            print(f"Ошибка очистки текста: {e}")
            return text

    def extract_text_from_images(self, image_paths: List[str]) -> List[str]:
        """Применяет OCR ко всем изображениям из списка с улучшенной обработкой текста."""
        texts = []
        
        for image_path in image_paths:
            try:
                # Предобработка изображения
                processed_image = self.preprocess_image(image_path)
                
                # Применяем OCR с поддержкой обоих языков
                config = r'--oem 3 --psm 3 -c preserve_interword_spaces=1'
                
                # Используем оба языка одновременно
                ocr_text = pytesseract.image_to_string(
                    processed_image,
                    lang="rus+eng",
                    config=config
                )
                
                # Очищаем и нормализуем текст
                cleaned_text = self.clean_ocr_text(ocr_text)
                
                # Дополнительная проверка на качество распознавания
                if len(cleaned_text) > 0:
                    # Разбиваем на строки и фильтруем пустые
                    lines = [line.strip() for line in cleaned_text.split('\n') if line.strip()]
                    # Объединяем обратно в текст
                    final_text = '\n'.join(lines)
                    texts.append(final_text)
                else:
                    texts.append("")
                    
            except Exception as e:
                print(f"Ошибка OCR для {image_path}: {e}")
                texts.append("")
                
        return texts


    """
    Очистка текста
    """ 
    def decode_unicode_sequence(self, text: str) -> str:
        """Декодирует строки Unicode в читаемый текст."""
        # Найдем все /uniXXXX и преобразуем их в символы Unicode
        decoded_text = re.sub(
            r"/uni([0-9A-Fa-f]{4})",  # Паттерн для /uniXXXX
            lambda x: chr(int(x.group(1), 16)),  # Преобразование кодов в Unicode-символы
            text
        )
        return decoded_text

    def clean_text(self, text: str) -> str:
        """Очищает текст от лишних пробелов, некорректных символов и декодирует Unicode."""
        try:
            # Декодируем Unicode-последовательности
            text = self.decode_unicode_sequence(text)

            # Удаляем все специальные символы, оставляя буквы, цифры, точки и запятые
            text = re.sub(r'[^a-zA-Zа-яА-Я0-9.,\s]', '', text)

            # Удаляем лишние пробелы
            text = re.sub(r'\s+', ' ', text).strip()
            
            # Возвращаем очищенный текст
            return text
        except Exception as e:
            print(f"Ошибка очистки текста: {e}")
            return text

    """
    Генерация ID
    """
    def generate_document_id(self) -> str:
        """
        Генерация уникального ID документа с использованием uuid4.
        """
        return str(uuid.uuid4())

    """
    Векторизация
    """
    def vectorize_text(self, text: str) -> list:
        """Преобразует текст в dense vector."""
        try:
            dense_vector = self.vectorizer_model.encode(text, convert_to_numpy=True)
            return dense_vector.tolist()
        except Exception as e:
            print(f"Ошибка векторизации текста: {e}")
            return []

    """
    Процессинг документа из хранилища
    """
    def process_document(self, file_path: str) -> Dict[str, Any]:
        """Общий процесс обработки документа (PDF или Word)."""
        ext = os.path.splitext(file_path)[-1].lower()
        result = {}

        # Получаем название файла без расширения и создаем ID документа
        file_name = os.path.splitext(os.path.basename(file_path))[0]
        document_id = self.generate_document_id()    

        # Создаем папку для изображений документа
        folder_path = self.create_document_folder(document_id)

        # Извлекаем текст, метаданные и изображения
        if ext == ".pdf":
            text = self.extract_text_from_pdf(file_path)
            metadata = self.extract_metadata_pdf(file_path)
            images = self.extract_images_from_pdf(file_path, folder_path)
        elif ext == ".docx":
            text = self.extract_text_from_word(file_path)
            metadata = self.extract_metadata_word(file_path)
            images = self.extract_images_from_word(file_path, folder_path)
        else:
            print(f"Формат файла {ext} не поддерживается.")
            return {}

        # Применяем OCR для изображений
        ocr_texts = self.extract_text_from_images(images)

        # Очищаем текст
        cleaned_text = self.clean_text(text + " " + " ".join(ocr_texts))

        # Векторизуем очищенный текст
        text_vector = self.vectorize_text(cleaned_text)

        # Теги

        # Собираем итоговый результат
        result["document_id"] = document_id
        result["title"] = file_name  # Название файла без расширения
        result["text_content"] = cleaned_text
        result["text_content_vector"] = text_vector  # Добавляем векторное представление текста
        metadata["tags"] = self.generate_tags_multilang(cleaned_text, num_tags=10)  # Сгенерированные ключевые слова для текста
        result["metadata"] = metadata  # Все метаданные, включая теги

        # Формируем информацию о каждом изображении
        result["images"] = []
        for idx, (img_path, ocr_text) in enumerate(zip(images, ocr_texts), start=1):
            image_info = {
                "image_id": f"img_{idx}",
                "ocr_text": ocr_text,
                "position": f"Page {idx}",
                "image_path": img_path
            }
            result["images"].append(image_info)

        return result

@lru_cache()
def get_preprocessing_service():
    return PreprocessingService(
        stopwords_collection=STOPWORD_COLLECTION,
        vectorizer_model=VECTORIZER_MODEL
    )























